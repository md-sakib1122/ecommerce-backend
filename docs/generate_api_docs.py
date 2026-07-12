"""Generate ``docs/API_Documentation.docx`` — the full API reference for the
E-Commerce Ordering & Payment backend.

This script is the *source* for the Word document: the ``.docx`` is a binary
that cannot be code-reviewed, so the human-readable content lives here and the
document is regenerated from it.

Usage
-----
    pip install python-docx        # one-off (also pinned in requirements.txt)
    python docs/generate_api_docs.py

The content is transcribed from the implemented API surface
(``app/api/v1/*``, ``app/schemas/*``, ``app/api/deps.py``, ``app/main.py``,
``app/core/config.py``). Regenerate whenever the API changes.
"""

from __future__ import annotations

import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# ── Palette ──────────────────────────────────────────────────────────────────
ACCENT = "1F4E79"          # table header fill (dark blue)
CODE_FILL = "F5F5F5"       # code block background
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY = RGBColor(0x60, 0x60, 0x60)

METHOD_COLORS = {
    "GET": RGBColor(0x2E, 0x7D, 0x32),      # green
    "POST": RGBColor(0x15, 0x65, 0xC0),     # blue
    "PATCH": RGBColor(0xE6, 0x51, 0x00),    # orange
    "PUT": RGBColor(0x6A, 0x1B, 0x9A),      # purple
    "DELETE": RGBColor(0xC6, 0x28, 0x28),   # red
}

OUTPUT = Path(__file__).resolve().parent / "API_Documentation.docx"


# ── Low-level docx helpers ───────────────────────────────────────────────────
def set_mono(run, size: int = 9) -> None:
    """Make a run monospace (Consolas)."""
    run.font.name = "Consolas"
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), "Consolas")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_code_block(doc, text: str, size: int = 9):
    """A light-grey, monospace block that preserves line breaks."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(8)
    p_pr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), CODE_FILL)
    p_pr.append(shd)
    lines = text.split("\n")
    for i, line in enumerate(lines):
        run = para.add_run(line)
        set_mono(run, size)
        if i != len(lines) - 1:
            run.add_break()
    return para


def make_table(doc, headers: list[str], rows: list[list[str]]):
    """A bordered table with an accent-shaded, white, bold header row."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        shade_cell(hdr_cells[i], ACCENT)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(value))
            run.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_page_break(doc) -> None:
    doc.add_page_break()


def label_line(doc, label: str, value: str) -> None:
    """A 'Label: value' paragraph with the label in bold."""
    para = doc.add_paragraph()
    run = para.add_run(f"{label}: ")
    run.bold = True
    para.add_run(value)


def add_toc(doc) -> None:
    """Insert a Word TOC field (auto-populates on open — see updateFields)."""
    para = doc.add_paragraph()
    run = para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click here and choose 'Update Field' to build the table of contents."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)


def enable_update_fields(doc) -> None:
    """Ask Word to refresh fields (the TOC) when the document is opened."""
    settings = doc.settings.element
    el = OxmlElement("w:updateFields")
    el.set(qn("w:val"), "true")
    settings.append(el)


def add_page_numbers(doc) -> None:
    footer = doc.sections[0].footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run("Page ").font.size = Pt(8)
    run = para.add_run()
    run.font.size = Pt(8)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def endpoint_heading(doc, method: str, path: str) -> None:
    """Heading 2 with a colour-coded method label."""
    heading = doc.add_heading("", level=2)
    m_run = heading.add_run(f"{method}  ")
    m_run.font.color.rgb = METHOD_COLORS.get(method, RGBColor(0, 0, 0))
    p_run = heading.add_run(path)
    p_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


# ── Content rendering ────────────────────────────────────────────────────────
def render_endpoint(doc, ep: dict) -> None:
    endpoint_heading(doc, ep["method"], ep["path"])
    label_line(doc, "Authentication", ep["auth"])
    doc.add_paragraph(ep["summary"])

    if ep.get("path_params"):
        doc.add_paragraph().add_run("Path parameters").bold = True
        make_table(doc, ["Name", "Type", "Description"], ep["path_params"])

    if ep.get("query_params"):
        doc.add_paragraph().add_run("Query parameters").bold = True
        make_table(doc, ["Name", "Type", "Default", "Description"], ep["query_params"])

    if ep.get("body_note"):
        doc.add_paragraph().add_run("Request body").bold = True
        doc.add_paragraph(ep["body_note"])
    if ep.get("body"):
        b = ep["body"]
        doc.add_paragraph().add_run("Request body").bold = True
        para = doc.add_paragraph()
        para.add_run("Schema: ").bold = True
        set_mono(para.add_run(b["schema"]))
        para.add_run("  (application/json)")
        make_table(
            doc,
            ["Field", "Type", "Required", "Constraints / Default", "Description"],
            b["fields"],
        )

    doc.add_paragraph().add_run("Responses").bold = True
    make_table(doc, ["Status", "Body", "Description"], ep["responses"])

    if ep.get("curl"):
        doc.add_paragraph().add_run("Example request").bold = True
        add_code_block(doc, ep["curl"])
    if ep.get("response_example"):
        doc.add_paragraph().add_run("Example response").bold = True
        add_code_block(doc, ep["response_example"])


def render_resource(doc, resource: dict) -> None:
    add_page_break(doc)
    doc.add_heading(resource["title"], level=1)
    if resource.get("intro"):
        doc.add_paragraph(resource["intro"])
    for ep in resource["endpoints"]:
        render_endpoint(doc, ep)


# ═════════════════════════════════════════════════════════════════════════════
# DATA
# ═════════════════════════════════════════════════════════════════════════════
ENUMS = [
    ["ProductStatus", "active, inactive", "app/models/product.py"],
    ["OrderStatus", "pending, paid, canceled", "app/models/order.py"],
    ["PaymentProvider", "stripe, bkash", "app/models/payment.py"],
    ["PaymentStatus", "pending, success, failed", "app/models/payment.py"],
]

STATUS_CODES = [
    ["200 OK", "Successful read/update, or an idempotent no-op."],
    ["201 Created", "Resource created (register, create product/category/order)."],
    ["204 No Content", "Successful delete — empty body."],
    ["400 Bad Request", "Invalid state or mismatched identifiers (e.g. inactive product, body/URL id mismatch, invalid webhook signature)."],
    ["401 Unauthorized", "Missing / invalid / expired bearer token, or bad login credentials."],
    ["403 Forbidden", "Authenticated but not an admin on an admin-only route."],
    ["404 Not Found", "Resource does not exist, or is not owned by the caller."],
    ["409 Conflict", "Duplicate unique value (email, SKU, transaction_id) or an illegal state transition."],
    ["422 Unprocessable Entity", "Request body failed Pydantic validation (FastAPI default)."],
    ["502 Bad Gateway", "Upstream payment provider (Stripe / bKash) returned an unusable response."],
]

RESOURCES = [
    # ── AUTH ─────────────────────────────────────────────────────────────────
    {
        "title": "Authentication",
        "intro": "Public endpoints for creating an account and obtaining a bearer "
                 "token. All other write endpoints require the token returned by "
                 "login.",
        "endpoints": [
            {
                "method": "POST",
                "path": "/api/v1/auth/register",
                "auth": "Public",
                "summary": "Register a new (non-admin) user account. The email must be unique; "
                           "the password is hashed server-side with bcrypt and never stored or returned in plaintext.",
                "body": {
                    "schema": "UserCreate",
                    "fields": [
                        ["email", "string (email)", "Yes", "RFC-validated", "Login identifier; must be unique."],
                        ["password", "string", "Yes", "min 8, max 128", "Plaintext password (hashed on the server)."],
                        ["full_name", "string | null", "No", "default null", "Optional display name."],
                    ],
                },
                "responses": [
                    ["201 Created", "UserRead", "The created user (no password field)."],
                    ["409 Conflict", "{ detail }", "Email already registered."],
                    ["422", "{ detail }", "Validation error (e.g. password too short)."],
                ],
                "curl": "curl -X POST http://localhost:8000/api/v1/auth/register \\\n"
                        "  -H \"Content-Type: application/json\" \\\n"
                        "  -d '{\"email\":\"jane@example.com\",\"password\":\"s3cretpassword\",\"full_name\":\"Jane Doe\"}'",
                "response_example": "{\n"
                                    "  \"id\": 1,\n"
                                    "  \"email\": \"jane@example.com\",\n"
                                    "  \"full_name\": \"Jane Doe\",\n"
                                    "  \"is_admin\": false,\n"
                                    "  \"is_active\": true,\n"
                                    "  \"created_at\": \"2026-07-09T12:00:00Z\"\n"
                                    "}",
            },
            {
                "method": "POST",
                "path": "/api/v1/auth/login",
                "auth": "Public",
                "summary": "Exchange email + password for a bearer JWT. NOTE: this endpoint accepts a "
                           "JSON body (not OAuth2 form-encoding), so the Swagger 'Authorize' button will "
                           "not work as-is — send the JSON body shown below and use the returned token in "
                           "the Authorization header.",
                "body": {
                    "schema": "UserLogin",
                    "fields": [
                        ["email", "string (email)", "Yes", "—", "Registered email."],
                        ["password", "string", "Yes", "—", "Account password."],
                    ],
                },
                "responses": [
                    ["200 OK", "Token", "access_token (JWT, HS256) + token_type='bearer'. Token is valid for 24 hours."],
                    ["401 Unauthorized", "{ detail }", "Incorrect email or password (also fails for deactivated accounts)."],
                ],
                "curl": "curl -X POST http://localhost:8000/api/v1/auth/login \\\n"
                        "  -H \"Content-Type: application/json\" \\\n"
                        "  -d '{\"email\":\"jane@example.com\",\"password\":\"s3cretpassword\"}'",
                "response_example": "{\n"
                                    "  \"access_token\": \"eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9...\",\n"
                                    "  \"token_type\": \"bearer\"\n"
                                    "}",
            },
        ],
    },
    # ── USERS ────────────────────────────────────────────────────────────────
    {
        "title": "Users",
        "intro": "Self-service '/me' endpoints for the authenticated caller, plus admin-only "
                 "user administration. Admin routes return 401 without a valid token and 403 for "
                 "a valid non-admin token.",
        "endpoints": [
            {
                "method": "GET",
                "path": "/api/v1/users/me",
                "auth": "Bearer token",
                "summary": "Return the authenticated user's own profile.",
                "responses": [["200 OK", "UserRead", "The current user."]],
                "curl": "curl http://localhost:8000/api/v1/users/me \\\n  -H \"Authorization: Bearer $TOKEN\"",
                "response_example": "{\n  \"id\": 1,\n  \"email\": \"jane@example.com\",\n  \"full_name\": \"Jane Doe\",\n"
                                    "  \"is_admin\": false,\n  \"is_active\": true,\n  \"created_at\": \"2026-07-09T12:00:00Z\"\n}",
            },
            {
                "method": "PATCH",
                "path": "/api/v1/users/me",
                "auth": "Bearer token",
                "summary": "Update your own profile. Only full_name is applied here; is_active is "
                           "ignored for self-service (changing it requires an admin route).",
                "body": {
                    "schema": "UserUpdate",
                    "fields": [
                        ["full_name", "string | null", "No", "default null", "New display name."],
                        ["is_active", "boolean | null", "No", "default null", "Ignored on this route."],
                    ],
                },
                "responses": [["200 OK", "UserRead", "The updated profile."]],
                "curl": "curl -X PATCH http://localhost:8000/api/v1/users/me \\\n"
                        "  -H \"Authorization: Bearer $TOKEN\" -H \"Content-Type: application/json\" \\\n"
                        "  -d '{\"full_name\":\"Jane A. Doe\"}'",
                "response_example": "{\n  \"id\": 1,\n  \"email\": \"jane@example.com\",\n  \"full_name\": \"Jane A. Doe\",\n"
                                    "  \"is_admin\": false,\n  \"is_active\": true,\n  \"created_at\": \"2026-07-09T12:00:00Z\"\n}",
            },
            {
                "method": "GET",
                "path": "/api/v1/users/me/orders",
                "auth": "Bearer token",
                "summary": "List the caller's own orders, newest first.",
                "responses": [["200 OK", "OrderRead[]", "Array of the caller's orders (each with nested items)."]],
                "curl": "curl http://localhost:8000/api/v1/users/me/orders \\\n  -H \"Authorization: Bearer $TOKEN\"",
                "response_example": "[\n  {\n    \"id\": 100,\n    \"user_id\": 1,\n    \"total_amount\": \"1998.00\",\n"
                                    "    \"status\": \"paid\",\n    \"items\": [ ... ],\n    \"created_at\": \"2026-07-09T12:05:00Z\"\n  }\n]",
            },
            {
                "method": "GET",
                "path": "/api/v1/users/me/payments",
                "auth": "Bearer token",
                "summary": "List the caller's own payments (resolved through their orders).",
                "responses": [["200 OK", "PaymentRead[]", "Array of the caller's payment records."]],
                "curl": "curl http://localhost:8000/api/v1/users/me/payments \\\n  -H \"Authorization: Bearer $TOKEN\"",
                "response_example": "[\n  {\n    \"id\": 1,\n    \"order_id\": 100,\n    \"provider\": \"stripe\",\n"
                                    "    \"transaction_id\": \"pi_3Q...\",\n    \"status\": \"success\",\n"
                                    "    \"created_at\": \"2026-07-09T12:06:00Z\"\n  }\n]",
            },
            {
                "method": "GET",
                "path": "/api/v1/users",
                "auth": "Admin",
                "summary": "Paginated list of all users.",
                "query_params": [
                    ["skip", "integer", "0", "Number of records to skip (offset)."],
                    ["limit", "integer", "100", "Maximum records to return."],
                ],
                "responses": [
                    ["200 OK", "UserRead[]", "Page of users."],
                    ["401 / 403", "{ detail }", "Missing token / not an admin."],
                ],
                "curl": "curl \"http://localhost:8000/api/v1/users?skip=0&limit=50\" \\\n  -H \"Authorization: Bearer $ADMIN_TOKEN\"",
                "response_example": "[\n  { \"id\": 1, \"email\": \"jane@example.com\", \"full_name\": \"Jane Doe\",\n"
                                    "    \"is_admin\": false, \"is_active\": true, \"created_at\": \"2026-07-09T12:00:00Z\" }\n]",
            },
            {
                "method": "GET",
                "path": "/api/v1/users/{user_id}",
                "auth": "Admin",
                "summary": "Fetch a single user by id.",
                "path_params": [["user_id", "integer", "Target user id."]],
                "responses": [
                    ["200 OK", "UserRead", "The requested user."],
                    ["404 Not Found", "{ detail }", "No user with that id."],
                    ["401 / 403", "{ detail }", "Missing token / not an admin."],
                ],
                "curl": "curl http://localhost:8000/api/v1/users/1 \\\n  -H \"Authorization: Bearer $ADMIN_TOKEN\"",
                "response_example": "{\n  \"id\": 1,\n  \"email\": \"jane@example.com\",\n  \"full_name\": \"Jane Doe\",\n"
                                    "  \"is_admin\": false,\n  \"is_active\": true,\n  \"created_at\": \"2026-07-09T12:00:00Z\"\n}",
            },
            {
                "method": "PATCH",
                "path": "/api/v1/users/{user_id}",
                "auth": "Admin",
                "summary": "Admin activate / deactivate a user. Uses the is_active field (defaults to "
                           "true if omitted); full_name is ignored on this route. Deactivated users lose "
                           "access on their next request.",
                "path_params": [["user_id", "integer", "Target user id."]],
                "body": {
                    "schema": "UserUpdate",
                    "fields": [
                        ["is_active", "boolean | null", "No", "default true", "New active state."],
                        ["full_name", "string | null", "No", "—", "Ignored on this route."],
                    ],
                },
                "responses": [
                    ["200 OK", "UserRead", "The updated user."],
                    ["404 Not Found", "{ detail }", "No user with that id."],
                    ["401 / 403", "{ detail }", "Missing token / not an admin."],
                ],
                "curl": "curl -X PATCH http://localhost:8000/api/v1/users/2 \\\n"
                        "  -H \"Authorization: Bearer $ADMIN_TOKEN\" -H \"Content-Type: application/json\" \\\n"
                        "  -d '{\"is_active\":false}'",
                "response_example": "{\n  \"id\": 2,\n  \"email\": \"bob@example.com\",\n  \"full_name\": \"Bob\",\n"
                                    "  \"is_admin\": false,\n  \"is_active\": false,\n  \"created_at\": \"2026-07-08T09:00:00Z\"\n}",
            },
        ],
    },
    # ── CATEGORIES ───────────────────────────────────────────────────────────
    {
        "title": "Categories",
        "intro": "A self-referential category tree (adjacency list; parent_id = null means root). "
                 "Reads are public; writes are admin-only. The tree endpoint is a DFS traversal cached "
                 "in Redis and invalidated on any category write.",
        "endpoints": [
            {
                "method": "GET",
                "path": "/api/v1/categories",
                "auth": "Public",
                "summary": "Flat, paginated list of categories.",
                "query_params": [
                    ["skip", "integer", "0", "Offset."],
                    ["limit", "integer", "100", "Max records."],
                ],
                "responses": [["200 OK", "CategoryRead[]", "Page of categories."]],
                "curl": "curl \"http://localhost:8000/api/v1/categories?limit=50\"",
                "response_example": "[\n  { \"id\": 1, \"name\": \"Electronics\", \"parent_id\": null },\n"
                                    "  { \"id\": 2, \"name\": \"Phones\", \"parent_id\": 1 }\n]",
            },
            {
                "method": "GET",
                "path": "/api/v1/categories/tree",
                "auth": "Public",
                "summary": "Return the category tree as nested nodes (Redis-cached DFS). With no "
                           "root_id the full forest is returned; with root_id only that subtree.",
                "query_params": [
                    ["root_id", "integer | null", "null", "Subtree root; null returns the whole forest."],
                ],
                "responses": [
                    ["200 OK", "CategoryTreeNode[]", "Nested tree; each node has a recursive 'children' array."],
                    ["404 Not Found", "{ detail }", "root_id does not exist."],
                ],
                "curl": "curl \"http://localhost:8000/api/v1/categories/tree\"",
                "response_example": "[\n  {\n    \"id\": 1, \"name\": \"Electronics\", \"parent_id\": null,\n"
                                    "    \"children\": [\n      { \"id\": 2, \"name\": \"Phones\", \"parent_id\": 1, \"children\": [] }\n"
                                    "    ]\n  }\n]",
            },
            {
                "method": "GET",
                "path": "/api/v1/categories/{category_id}",
                "auth": "Public",
                "summary": "Fetch a single category.",
                "path_params": [["category_id", "integer", "Target category id."]],
                "responses": [
                    ["200 OK", "CategoryRead", "The category."],
                    ["404 Not Found", "{ detail }", "No category with that id."],
                ],
                "curl": "curl http://localhost:8000/api/v1/categories/2",
                "response_example": "{ \"id\": 2, \"name\": \"Phones\", \"parent_id\": 1 }",
            },
            {
                "method": "GET",
                "path": "/api/v1/categories/{category_id}/products",
                "auth": "Public",
                "summary": "List active products in this category and all of its descendants (DFS over the subtree).",
                "path_params": [["category_id", "integer", "Branch root category id."]],
                "query_params": [
                    ["skip", "integer", "0", "Offset."],
                    ["limit", "integer", "100", "Max records."],
                ],
                "responses": [
                    ["200 OK", "ProductRead[]", "Active products across the branch."],
                    ["404 Not Found", "{ detail }", "Category does not exist."],
                ],
                "curl": "curl \"http://localhost:8000/api/v1/categories/1/products?limit=20\"",
                "response_example": "[\n  { \"id\": 10, \"name\": \"iPhone 15\", \"sku\": \"IP15-128\", \"price\": \"999.00\",\n"
                                    "    \"stock\": 25, \"status\": \"active\", \"category_id\": 2, \"created_at\": \"...\" }\n]",
            },
            {
                "method": "POST",
                "path": "/api/v1/categories",
                "auth": "Admin",
                "summary": "Create a category. Provide parent_id to nest it, or omit / null for a root category.",
                "body": {
                    "schema": "CategoryCreate",
                    "fields": [
                        ["name", "string", "Yes", "max 150", "Category name."],
                        ["parent_id", "integer | null", "No", "default null", "Parent category (null = root)."],
                    ],
                },
                "responses": [
                    ["201 Created", "CategoryRead", "The created category."],
                    ["400 Bad Request", "{ detail }", "parent_id references a category that does not exist."],
                    ["401 / 403", "{ detail }", "Missing token / not an admin."],
                ],
                "curl": "curl -X POST http://localhost:8000/api/v1/categories \\\n"
                        "  -H \"Authorization: Bearer $ADMIN_TOKEN\" -H \"Content-Type: application/json\" \\\n"
                        "  -d '{\"name\":\"Laptops\",\"parent_id\":1}'",
                "response_example": "{ \"id\": 3, \"name\": \"Laptops\", \"parent_id\": 1 }",
            },
            {
                "method": "PATCH",
                "path": "/api/v1/categories/{category_id}",
                "auth": "Admin",
                "summary": "Partial update. Sending parent_id: null explicitly re-parents to root; omitting "
                           "parent_id leaves the parent unchanged. Cycles and self-parenting are rejected.",
                "path_params": [["category_id", "integer", "Target category id."]],
                "body": {
                    "schema": "CategoryUpdate",
                    "fields": [
                        ["name", "string | null", "No", "max 150", "New name."],
                        ["parent_id", "integer | null", "No", "—", "New parent (explicit null = move to root)."],
                    ],
                },
                "responses": [
                    ["200 OK", "CategoryRead", "The updated category."],
                    ["404 Not Found", "{ detail }", "Category does not exist."],
                    ["400 Bad Request", "{ detail }", "Self-parenting, missing parent, or would create a cycle."],
                    ["401 / 403", "{ detail }", "Missing token / not an admin."],
                ],
                "curl": "curl -X PATCH http://localhost:8000/api/v1/categories/3 \\\n"
                        "  -H \"Authorization: Bearer $ADMIN_TOKEN\" -H \"Content-Type: application/json\" \\\n"
                        "  -d '{\"parent_id\":null}'",
                "response_example": "{ \"id\": 3, \"name\": \"Laptops\", \"parent_id\": null }",
            },
            {
                "method": "DELETE",
                "path": "/api/v1/categories/{category_id}",
                "auth": "Admin",
                "summary": "Delete a category. Refused if it still has subcategories or products.",
                "path_params": [["category_id", "integer", "Target category id."]],
                "responses": [
                    ["204 No Content", "(empty)", "Deleted."],
                    ["404 Not Found", "{ detail }", "Category does not exist."],
                    ["409 Conflict", "{ detail }", "Category has subcategories or products."],
                    ["401 / 403", "{ detail }", "Missing token / not an admin."],
                ],
                "curl": "curl -X DELETE http://localhost:8000/api/v1/categories/3 \\\n  -H \"Authorization: Bearer $ADMIN_TOKEN\"",
            },
        ],
    },
    # ── PRODUCTS ─────────────────────────────────────────────────────────────
    {
        "title": "Products",
        "intro": "Catalog browsing is public and returns active products by default; product writes are "
                 "admin-only. Prices are DECIMAL(12,2) serialized as JSON strings. Delete is a soft delete "
                 "(status -> inactive).",
        "endpoints": [
            {
                "method": "GET",
                "path": "/api/v1/products",
                "auth": "Public",
                "summary": "Browse the catalog with optional filters and a name search.",
                "query_params": [
                    ["skip", "integer", "0", "Offset."],
                    ["limit", "integer", "100", "Max records."],
                    ["status", "ProductStatus", "active", "Filter by status (active / inactive)."],
                    ["category_id", "integer | null", "null", "Filter to a single category."],
                    ["search", "string | null", "null", "Case-insensitive substring match on name."],
                ],
                "responses": [["200 OK", "ProductRead[]", "Matching products."]],
                "curl": "curl \"http://localhost:8000/api/v1/products?search=iphone&limit=20\"",
                "response_example": "[\n  { \"id\": 10, \"name\": \"iPhone 15\", \"sku\": \"IP15-128\",\n"
                                    "    \"description\": \"128GB, black\", \"price\": \"999.00\", \"stock\": 25,\n"
                                    "    \"status\": \"active\", \"category_id\": 2, \"created_at\": \"2026-07-01T00:00:00Z\" }\n]",
            },
            {
                "method": "GET",
                "path": "/api/v1/products/{product_id}",
                "auth": "Public",
                "summary": "Fetch a single product.",
                "path_params": [["product_id", "integer", "Target product id."]],
                "responses": [
                    ["200 OK", "ProductRead", "The product."],
                    ["404 Not Found", "{ detail }", "No product with that id."],
                ],
                "curl": "curl http://localhost:8000/api/v1/products/10",
                "response_example": "{ \"id\": 10, \"name\": \"iPhone 15\", \"sku\": \"IP15-128\", \"description\": \"128GB, black\",\n"
                                    "  \"price\": \"999.00\", \"stock\": 25, \"status\": \"active\", \"category_id\": 2,\n"
                                    "  \"created_at\": \"2026-07-01T00:00:00Z\" }",
            },
            {
                "method": "GET",
                "path": "/api/v1/products/{product_id}/recommendations",
                "auth": "Public",
                "summary": "Related products, found by a DFS over the product's category branch. Returns an "
                           "empty array if the product has no category.",
                "path_params": [["product_id", "integer", "Anchor product id."]],
                "query_params": [["limit", "integer", "10", "Max recommendations."]],
                "responses": [
                    ["200 OK", "ProductRead[]", "Recommended products (may be empty)."],
                    ["404 Not Found", "{ detail }", "Anchor product does not exist."],
                ],
                "curl": "curl \"http://localhost:8000/api/v1/products/10/recommendations?limit=5\"",
                "response_example": "[\n  { \"id\": 11, \"name\": \"iPhone 15 Pro\", \"sku\": \"IP15P-256\", \"price\": \"1199.00\",\n"
                                    "    \"stock\": 12, \"status\": \"active\", \"category_id\": 2, \"created_at\": \"...\" }\n]",
            },
            {
                "method": "POST",
                "path": "/api/v1/products",
                "auth": "Admin",
                "summary": "Create a product. SKU must be unique.",
                "body": {
                    "schema": "ProductCreate",
                    "fields": [
                        ["name", "string", "Yes", "max 255", "Product name."],
                        ["sku", "string", "Yes", "max 64, unique", "Stock-keeping unit."],
                        ["description", "string | null", "No", "default null", "Optional description."],
                        ["price", "decimal (string)", "Yes", "> 0", "Unit price, DECIMAL(12,2)."],
                        ["stock", "integer", "No", ">= 0, default 0", "Initial stock quantity."],
                        ["status", "ProductStatus", "No", "default active", "active / inactive."],
                        ["category_id", "integer | null", "No", "default null", "Owning category."],
                    ],
                },
                "responses": [
                    ["201 Created", "ProductRead", "The created product."],
                    ["409 Conflict", "{ detail }", "Duplicate SKU or a DB constraint violation."],
                    ["401 / 403", "{ detail }", "Missing token / not an admin."],
                ],
                "curl": "curl -X POST http://localhost:8000/api/v1/products \\\n"
                        "  -H \"Authorization: Bearer $ADMIN_TOKEN\" -H \"Content-Type: application/json\" \\\n"
                        "  -d '{\"name\":\"iPhone 15\",\"sku\":\"IP15-128\",\"price\":\"999.00\",\"stock\":25,\"category_id\":2}'",
                "response_example": "{ \"id\": 10, \"name\": \"iPhone 15\", \"sku\": \"IP15-128\", \"description\": null,\n"
                                    "  \"price\": \"999.00\", \"stock\": 25, \"status\": \"active\", \"category_id\": 2,\n"
                                    "  \"created_at\": \"2026-07-09T12:10:00Z\" }",
            },
            {
                "method": "PATCH",
                "path": "/api/v1/products/{product_id}",
                "auth": "Admin",
                "summary": "Partial update. SKU is immutable (not accepted here). All fields optional.",
                "path_params": [["product_id", "integer", "Target product id."]],
                "body": {
                    "schema": "ProductUpdate",
                    "fields": [
                        ["name", "string | null", "No", "max 255", "New name."],
                        ["description", "string | null", "No", "—", "New description."],
                        ["price", "decimal (string) | null", "No", "> 0", "New unit price."],
                        ["stock", "integer | null", "No", ">= 0", "New stock quantity."],
                        ["status", "ProductStatus | null", "No", "—", "New status."],
                        ["category_id", "integer | null", "No", "—", "New category."],
                    ],
                },
                "responses": [
                    ["200 OK", "ProductRead", "The updated product."],
                    ["404 Not Found", "{ detail }", "Product does not exist."],
                    ["409 Conflict", "{ detail }", "Constraint violation."],
                    ["401 / 403", "{ detail }", "Missing token / not an admin."],
                ],
                "curl": "curl -X PATCH http://localhost:8000/api/v1/products/10 \\\n"
                        "  -H \"Authorization: Bearer $ADMIN_TOKEN\" -H \"Content-Type: application/json\" \\\n"
                        "  -d '{\"price\":\"899.00\",\"stock\":30}'",
                "response_example": "{ \"id\": 10, \"name\": \"iPhone 15\", \"sku\": \"IP15-128\", \"description\": null,\n"
                                    "  \"price\": \"899.00\", \"stock\": 30, \"status\": \"active\", \"category_id\": 2,\n"
                                    "  \"created_at\": \"2026-07-09T12:10:00Z\" }",
            },
            {
                "method": "DELETE",
                "path": "/api/v1/products/{product_id}",
                "auth": "Admin",
                "summary": "Soft delete — sets the product's status to inactive rather than removing the row "
                           "(preserves historical order references).",
                "path_params": [["product_id", "integer", "Target product id."]],
                "responses": [
                    ["204 No Content", "(empty)", "Product deactivated."],
                    ["404 Not Found", "{ detail }", "Product does not exist."],
                    ["401 / 403", "{ detail }", "Missing token / not an admin."],
                ],
                "curl": "curl -X DELETE http://localhost:8000/api/v1/products/10 \\\n  -H \"Authorization: Bearer $ADMIN_TOKEN\"",
            },
        ],
    },
    # ── ORDERS ───────────────────────────────────────────────────────────────
    {
        "title": "Orders",
        "intro": "Authenticated users place and manage their own orders. All money is computed "
                 "server-side: order_items.price is a snapshot of the product price at order time, "
                 "subtotal = price x quantity, and total_amount = sum of subtotals. Client-supplied "
                 "prices are never trusted. Orders start as 'pending'; stock is only reduced after a "
                 "successful payment.",
        "endpoints": [
            {
                "method": "POST",
                "path": "/api/v1/orders",
                "auth": "Bearer token",
                "summary": "Place an order from a list of product/quantity lines. Totals are computed "
                           "server-side and the order starts in 'pending' status; no stock is reserved yet.",
                "body": {
                    "schema": "OrderCreate",
                    "fields": [
                        ["items", "OrderItemCreate[]", "Yes", "min 1 item", "Line items to order."],
                        ["items[].product_id", "integer", "Yes", "—", "Product to order."],
                        ["items[].quantity", "integer", "Yes", "> 0", "Quantity of that product."],
                    ],
                },
                "responses": [
                    ["201 Created", "OrderRead", "The created order with computed totals and items."],
                    ["404 Not Found", "{ detail }", "A referenced product does not exist."],
                    ["400 Bad Request", "{ detail }", "Product is inactive, or stock is insufficient."],
                ],
                "curl": "curl -X POST http://localhost:8000/api/v1/orders \\\n"
                        "  -H \"Authorization: Bearer $TOKEN\" -H \"Content-Type: application/json\" \\\n"
                        "  -d '{\"items\":[{\"product_id\":10,\"quantity\":2}]}'",
                "response_example": "{\n  \"id\": 100,\n  \"user_id\": 1,\n  \"total_amount\": \"1998.00\",\n"
                                    "  \"status\": \"pending\",\n  \"items\": [\n"
                                    "    { \"id\": 1, \"product_id\": 10, \"quantity\": 2, \"price\": \"999.00\", \"subtotal\": \"1998.00\" }\n"
                                    "  ],\n  \"created_at\": \"2026-07-09T12:15:00Z\"\n}",
            },
            {
                "method": "GET",
                "path": "/api/v1/orders/{order_id}",
                "auth": "Bearer token (owner)",
                "summary": "Fetch one of the caller's own orders. Requesting another user's order returns "
                           "404 (not 403), so ownership is not leaked.",
                "path_params": [["order_id", "integer", "Target order id."]],
                "responses": [
                    ["200 OK", "OrderRead", "The order."],
                    ["404 Not Found", "{ detail }", "Order does not exist or is not owned by the caller."],
                ],
                "curl": "curl http://localhost:8000/api/v1/orders/100 \\\n  -H \"Authorization: Bearer $TOKEN\"",
                "response_example": "{\n  \"id\": 100,\n  \"user_id\": 1,\n  \"total_amount\": \"1998.00\",\n"
                                    "  \"status\": \"pending\",\n  \"items\": [ ... ],\n  \"created_at\": \"2026-07-09T12:15:00Z\"\n}",
            },
            {
                "method": "POST",
                "path": "/api/v1/orders/{order_id}/cancel",
                "auth": "Bearer token (owner)",
                "summary": "Cancel your own order while it is still 'pending'. Idempotent if the order is "
                           "already canceled.",
                "path_params": [["order_id", "integer", "Target order id."]],
                "responses": [
                    ["200 OK", "OrderRead", "The order, now 'canceled'."],
                    ["404 Not Found", "{ detail }", "Order does not exist or is not owned by the caller."],
                    ["409 Conflict", "{ detail }", "Only pending orders can be canceled (e.g. already paid)."],
                ],
                "curl": "curl -X POST http://localhost:8000/api/v1/orders/100/cancel \\\n  -H \"Authorization: Bearer $TOKEN\"",
                "response_example": "{\n  \"id\": 100,\n  \"user_id\": 1,\n  \"total_amount\": \"1998.00\",\n"
                                    "  \"status\": \"canceled\",\n  \"items\": [ ... ],\n  \"created_at\": \"2026-07-09T12:15:00Z\"\n}",
            },
        ],
    },
    # ── PAYMENTS ─────────────────────────────────────────────────────────────
    {
        "title": "Payments",
        "intro": "Payments use the Strategy pattern: a provider factory selects Stripe or bKash, and order "
                 "logic never branches on the provider. transaction_id is the idempotency key (unique) — "
                 "webhook/callback replays are no-ops, so stock never double-decrements. The webhook and "
                 "callback endpoints are the only unauthenticated POST routes; they are trusted via Stripe "
                 "signature verification and bKash server-side execution respectively (a client-supplied "
                 "'status' is never trusted).",
        "endpoints": [
            {
                "method": "POST",
                "path": "/api/v1/orders/{order_id}/checkout",
                "auth": "Bearer token (owner)",
                "summary": "Initiate payment for one of your own 'pending' orders with the chosen provider. "
                           "Returns provider-specific data the client uses to complete payment (Stripe "
                           "client_secret, or a bKash redirect URL). The order_id in the body must match the URL.",
                "path_params": [["order_id", "integer", "Order to pay for (must match body.order_id)."]],
                "body": {
                    "schema": "PaymentInitiate",
                    "fields": [
                        ["order_id", "integer", "Yes", "must equal URL id", "Order to pay for."],
                        ["provider", "PaymentProvider", "Yes", "stripe | bkash", "Payment provider."],
                    ],
                },
                "responses": [
                    ["200 OK", "StripeCheckoutResponse | BkashCheckoutResponse",
                     "Stripe: { client_secret, payment_intent_id }. bKash: { bkash_url, payment_id }."],
                    ["400 Bad Request", "{ detail }", "Body order_id does not match the URL, or unsupported provider."],
                    ["404 Not Found", "{ detail }", "Order does not exist or is not owned by the caller."],
                    ["409 Conflict", "{ detail }", "Order is not 'pending', or duplicate transaction_id."],
                    ["502 Bad Gateway", "{ detail }", "bKash returned an unusable create response."],
                ],
                "curl": "curl -X POST http://localhost:8000/api/v1/orders/100/checkout \\\n"
                        "  -H \"Authorization: Bearer $TOKEN\" -H \"Content-Type: application/json\" \\\n"
                        "  -d '{\"order_id\":100,\"provider\":\"stripe\"}'",
                "response_example": "// Stripe\n{\n  \"client_secret\": \"pi_3Q..._secret_abc123\",\n"
                                    "  \"payment_intent_id\": \"pi_3Q...\"\n}\n\n"
                                    "// bKash\n{\n  \"bkash_url\": \"https://tokenized.sandbox.bka.sh/...\",\n"
                                    "  \"payment_id\": \"TR0011...\"\n}",
            },
            {
                "method": "POST",
                "path": "/api/v1/payments/stripe/webhook",
                "auth": "Public (Stripe signature-verified)",
                "summary": "Stripe webhook receiver. The raw request body is verified against the "
                           "Stripe-Signature header using STRIPE_WEBHOOK_SECRET. On 'payment_intent.succeeded' "
                           "the matching payment is settled and stock is reduced atomically; replays are no-ops.",
                "body_note": "Raw Stripe event JSON (no fixed Pydantic model). Requires the 'Stripe-Signature' "
                             "header. Do not construct this by hand — Stripe sends it. Only "
                             "'payment_intent.succeeded' and 'payment_intent.payment_failed' change payment state.",
                "responses": [
                    ["200 OK", "{ \"received\": true }", "Event accepted (including ignored event types)."],
                    ["400 Bad Request", "{ detail }", "Invalid or missing Stripe signature."],
                    ["404 Not Found", "{ detail }", "No local payment matches the event's transaction."],
                ],
                "curl": "# Sent by Stripe, not by clients. Locally you can forward events with the Stripe CLI:\n"
                        "stripe listen --forward-to localhost:8000/api/v1/payments/stripe/webhook",
                "response_example": "{ \"received\": true }",
            },
            {
                "method": "POST",
                "path": "/api/v1/payments/bkash/callback",
                "auth": "Public (server-side execute)",
                "summary": "bKash redirect callback. Only paymentID is used: the server calls bKash "
                           "'execute' to authoritatively settle the payment (the client-supplied 'status' is "
                           "ignored). On success the order is marked paid and stock is reduced; replays are no-ops.",
                "body": {
                    "schema": "BkashWebhookPayload",
                    "fields": [
                        ["paymentID", "string", "Yes", "—", "bKash payment id (the only field used)."],
                        ["status", "string", "No", "—", "Client-reported status (ignored / not trusted)."],
                        ["trxID", "string | null", "No", "default null", "bKash transaction id (informational)."],
                    ],
                },
                "responses": [
                    ["200 OK", "PaymentRead", "The settled payment record."],
                    ["400 Bad Request", "{ detail }", "Missing paymentID."],
                    ["404 Not Found", "{ detail }", "No local payment matches the paymentID."],
                    ["502 Bad Gateway", "{ detail }", "bKash execute call failed."],
                ],
                "curl": "curl -X POST http://localhost:8000/api/v1/payments/bkash/callback \\\n"
                        "  -H \"Content-Type: application/json\" \\\n"
                        "  -d '{\"paymentID\":\"TR0011...\",\"status\":\"success\",\"trxID\":\"BFD70ABCD\"}'",
                "response_example": "{\n  \"id\": 2,\n  \"order_id\": 100,\n  \"provider\": \"bkash\",\n"
                                    "  \"transaction_id\": \"TR0011...\",\n  \"status\": \"success\",\n"
                                    "  \"created_at\": \"2026-07-09T12:20:00Z\"\n}",
            },
            {
                "method": "GET",
                "path": "/api/v1/payments/{payment_id}",
                "auth": "Bearer token (owner)",
                "summary": "Fetch one of the caller's own payments (ownership resolved through the order).",
                "path_params": [["payment_id", "integer", "Target payment id."]],
                "responses": [
                    ["200 OK", "PaymentRead", "The payment."],
                    ["404 Not Found", "{ detail }", "Payment does not exist or is not owned by the caller."],
                ],
                "curl": "curl http://localhost:8000/api/v1/payments/1 \\\n  -H \"Authorization: Bearer $TOKEN\"",
                "response_example": "{\n  \"id\": 1,\n  \"order_id\": 100,\n  \"provider\": \"stripe\",\n"
                                    "  \"transaction_id\": \"pi_3Q...\",\n  \"status\": \"success\",\n"
                                    "  \"created_at\": \"2026-07-09T12:06:00Z\"\n}",
            },
        ],
    },
    # ── HEALTH ───────────────────────────────────────────────────────────────
    {
        "title": "Health",
        "intro": "Operational endpoint, mounted at the root (not under /api/v1) and hidden from the OpenAPI schema.",
        "endpoints": [
            {
                "method": "GET",
                "path": "/health",
                "auth": "Public",
                "summary": "Liveness/readiness probe. Runs 'SELECT 1' against the database and reports the environment.",
                "responses": [
                    ["200 OK", "{ status, env }", "Service and DB are reachable."],
                ],
                "curl": "curl http://localhost:8000/health",
                "response_example": "{ \"status\": \"ok\", \"env\": \"development\" }",
            },
        ],
    },
]

# ── Data-model appendix ──────────────────────────────────────────────────────
SCHEMA_APPENDIX = [
    ("UserCreate  (request)", [
        ["email", "string (email)", "Yes", "RFC-validated", "Login identifier; unique."],
        ["password", "string", "Yes", "min 8, max 128", "Plaintext password (hashed server-side)."],
        ["full_name", "string | null", "No", "default null", "Display name."],
    ]),
    ("UserLogin  (request)", [
        ["email", "string (email)", "Yes", "—", "Registered email."],
        ["password", "string", "Yes", "—", "Account password."],
    ]),
    ("UserUpdate  (request)", [
        ["full_name", "string | null", "No", "default null", "New display name."],
        ["is_active", "boolean | null", "No", "default null", "Active flag (admin route only)."],
    ]),
    ("UserRead  (response)", [
        ["id", "integer", "—", "—", "User id."],
        ["email", "string (email)", "—", "—", "Email."],
        ["full_name", "string | null", "—", "—", "Display name."],
        ["is_admin", "boolean", "—", "—", "Admin flag."],
        ["is_active", "boolean", "—", "—", "Active flag."],
        ["created_at", "datetime", "—", "—", "Creation timestamp."],
    ]),
    ("Token  (response)", [
        ["access_token", "string", "—", "—", "JWT bearer token (HS256)."],
        ["token_type", "string", "—", "default 'bearer'", "Token scheme."],
    ]),
    ("ProductCreate  (request)", [
        ["name", "string", "Yes", "max 255", "Product name."],
        ["sku", "string", "Yes", "max 64, unique", "Stock-keeping unit."],
        ["description", "string | null", "No", "default null", "Description."],
        ["price", "decimal (string)", "Yes", "> 0", "Unit price, DECIMAL(12,2)."],
        ["stock", "integer", "No", ">= 0, default 0", "Stock quantity."],
        ["status", "ProductStatus", "No", "default active", "active / inactive."],
        ["category_id", "integer | null", "No", "default null", "Owning category."],
    ]),
    ("ProductUpdate  (request, all optional)", [
        ["name", "string | null", "No", "max 255", "New name."],
        ["description", "string | null", "No", "—", "New description."],
        ["price", "decimal (string) | null", "No", "> 0", "New price."],
        ["stock", "integer | null", "No", ">= 0", "New stock."],
        ["status", "ProductStatus | null", "No", "—", "New status."],
        ["category_id", "integer | null", "No", "—", "New category."],
    ]),
    ("ProductRead  (response)", [
        ["id", "integer", "—", "—", "Product id."],
        ["name", "string", "—", "—", "Name."],
        ["sku", "string", "—", "—", "SKU."],
        ["description", "string | null", "—", "—", "Description."],
        ["price", "decimal (string)", "—", "—", "Unit price."],
        ["stock", "integer", "—", "—", "Stock quantity."],
        ["status", "ProductStatus", "—", "—", "active / inactive."],
        ["category_id", "integer | null", "—", "—", "Owning category."],
        ["created_at", "datetime", "—", "—", "Creation timestamp."],
    ]),
    ("CategoryCreate  (request)", [
        ["name", "string", "Yes", "max 150", "Category name."],
        ["parent_id", "integer | null", "No", "default null", "Parent (null = root)."],
    ]),
    ("CategoryUpdate  (request, all optional)", [
        ["name", "string | null", "No", "max 150", "New name."],
        ["parent_id", "integer | null", "No", "—", "New parent (explicit null = root)."],
    ]),
    ("CategoryRead  (response)", [
        ["id", "integer", "—", "—", "Category id."],
        ["name", "string", "—", "—", "Name."],
        ["parent_id", "integer | null", "—", "—", "Parent id (null = root)."],
    ]),
    ("CategoryTreeNode  (response, extends CategoryRead)", [
        ["id", "integer", "—", "—", "Category id."],
        ["name", "string", "—", "—", "Name."],
        ["parent_id", "integer | null", "—", "—", "Parent id."],
        ["children", "CategoryTreeNode[]", "—", "default []", "Recursive child nodes."],
    ]),
    ("OrderItemCreate  (request, nested)", [
        ["product_id", "integer", "Yes", "—", "Product to order."],
        ["quantity", "integer", "Yes", "> 0", "Quantity."],
    ]),
    ("OrderCreate  (request)", [
        ["items", "OrderItemCreate[]", "Yes", "min 1", "Line items."],
    ]),
    ("OrderItemRead  (response, nested)", [
        ["id", "integer", "—", "—", "Line-item id."],
        ["product_id", "integer", "—", "—", "Product id."],
        ["quantity", "integer", "—", "—", "Quantity."],
        ["price", "decimal (string)", "—", "—", "Price snapshot at order time."],
        ["subtotal", "decimal (string)", "—", "—", "price x quantity."],
    ]),
    ("OrderRead  (response)", [
        ["id", "integer", "—", "—", "Order id."],
        ["user_id", "integer", "—", "—", "Owner user id."],
        ["total_amount", "decimal (string)", "—", "—", "Sum of item subtotals."],
        ["status", "OrderStatus", "—", "—", "pending / paid / canceled."],
        ["items", "OrderItemRead[]", "—", "—", "Line items."],
        ["created_at", "datetime", "—", "—", "Creation timestamp."],
    ]),
    ("PaymentInitiate  (request)", [
        ["order_id", "integer", "Yes", "must match URL", "Order to pay for."],
        ["provider", "PaymentProvider", "Yes", "stripe | bkash", "Payment provider."],
    ]),
    ("PaymentRead  (response)", [
        ["id", "integer", "—", "—", "Payment id."],
        ["order_id", "integer", "—", "—", "Related order."],
        ["provider", "PaymentProvider", "—", "—", "stripe / bkash."],
        ["transaction_id", "string", "—", "unique", "Provider transaction id (idempotency key)."],
        ["status", "PaymentStatus", "—", "—", "pending / success / failed."],
        ["created_at", "datetime", "—", "—", "Creation timestamp."],
    ]),
    ("StripeCheckoutResponse  (response)", [
        ["client_secret", "string", "—", "—", "Stripe PaymentIntent client secret."],
        ["payment_intent_id", "string", "—", "—", "Stripe PaymentIntent id."],
    ]),
    ("BkashCheckoutResponse  (response)", [
        ["bkash_url", "string", "—", "—", "Redirect URL to complete payment."],
        ["payment_id", "string", "—", "—", "bKash payment id."],
    ]),
    ("StripeWebhookPayload  (schema; raw event used at runtime)", [
        ["id", "string", "—", "—", "Event id."],
        ["type", "string", "—", "—", "Event type."],
        ["data", "object", "—", "—", "Event data payload."],
    ]),
    ("BkashWebhookPayload  (request — callback)", [
        ["paymentID", "string", "Yes", "—", "bKash payment id (only field used)."],
        ["status", "string", "No", "—", "Client-reported status (ignored)."],
        ["trxID", "string | null", "No", "default null", "bKash transaction id."],
    ]),
]


# ═════════════════════════════════════════════════════════════════════════════
# BUILD
# ═════════════════════════════════════════════════════════════════════════════
def build() -> None:
    doc = Document()

    # Base font
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # ── Title page ───────────────────────────────────────────────────────────
    title = doc.add_heading("", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("E-Commerce Ordering & Payment API")
    run.font.size = Pt(26)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = subtitle.add_run("API Reference Documentation")
    s_run.font.size = Pt(15)
    s_run.font.color.rgb = GREY
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m_run = meta.add_run(
        f"Version 0.1.0   ·   Generated {datetime.date.today().isoformat()}   ·   Base URL: http://localhost:8000"
    )
    m_run.font.size = Pt(10)
    m_run.font.color.rgb = GREY

    doc.add_paragraph()
    doc.add_heading("Contents", level=1)
    add_toc(doc)

    # ── Overview ─────────────────────────────────────────────────────────────
    add_page_break(doc)
    doc.add_heading("Overview", level=1)
    doc.add_paragraph(
        "This is the REST API for an e-commerce ordering and payment system: it manages users, a "
        "category tree, a product catalog, orders, and payments across two providers (Stripe and bKash). "
        "It is built with FastAPI and SQLModel over PostgreSQL, uses Redis to cache the category tree, "
        "and integrates payments behind a provider-agnostic Strategy interface."
    )
    doc.add_paragraph(
        "Interactive documentation generated by FastAPI is available at /docs (Swagger UI) and /redoc "
        "while the server runs in a non-production environment. This document is a stable, offline "
        "companion to that live schema."
    )

    # ── Conventions ──────────────────────────────────────────────────────────
    doc.add_heading("Conventions", level=1)
    doc.add_heading("Base URL & versioning", level=2)
    doc.add_paragraph(
        "All resource endpoints are served under the /api/v1 prefix (e.g. /api/v1/products). "
        "The operational /health endpoint is the one exception — it lives at the root and is hidden "
        "from the OpenAPI schema."
    )
    doc.add_heading("Content type", level=2)
    doc.add_paragraph(
        "Requests and responses are JSON (application/json), except the Stripe webhook, which receives a "
        "raw signed body. Response bodies are mapped from ORM objects onto dedicated read schemas — internal "
        "fields such as password hashes and raw provider responses are never exposed."
    )
    doc.add_heading("Money", level=2)
    doc.add_paragraph(
        "All monetary values are DECIMAL(12,2) and are serialized as JSON strings (e.g. \"999.00\") to avoid "
        "binary floating-point rounding. Never send prices or totals from the client — they are always "
        "computed server-side."
    )
    doc.add_heading("Pagination", level=2)
    doc.add_paragraph(
        "List endpoints accept skip (offset, default 0) and limit (default 100) query parameters. "
        "The recommendations endpoint uses limit with a default of 10."
    )
    doc.add_heading("Enumerations", level=2)
    make_table(doc, ["Enum", "Values", "Defined in"], ENUMS)

    # ── Authentication ───────────────────────────────────────────────────────
    add_page_break(doc)
    doc.add_heading("Authentication & Authorization", level=1)
    doc.add_paragraph(
        "The API uses stateless bearer JWTs (HS256). Register an account, log in to receive a token, then "
        "send it on every protected request:"
    )
    add_code_block(doc, "Authorization: Bearer <access_token>")
    doc.add_paragraph(
        "Tokens carry the claims sub (user id), exp (expiry) and type ('access'), and are valid for 24 hours "
        "(1440 minutes). The user is re-loaded from the database on every request, so deactivating a user "
        "revokes access immediately."
    )
    doc.add_heading("Authorization model", level=2)
    doc.add_paragraph(
        "Authorization is a boolean is_admin flag (there are no string roles). Protected routes return 401 "
        "when the token is missing, malformed, expired, or the user is inactive; admin-only routes "
        "additionally return 403 for a valid non-admin token."
    )
    doc.add_heading("Important: login is JSON, not OAuth2 form data", level=2)
    doc.add_paragraph(
        "Although Swagger advertises an OAuth2 password flow, POST /api/v1/auth/login actually expects a "
        "JSON body ({\"email\": ..., \"password\": ...}). The Swagger 'Authorize' button (which posts "
        "form-encoded credentials) will therefore not work as-is; call the login endpoint directly and paste "
        "the returned token into the Authorize dialog, or send the header manually."
    )

    # ── Error handling ───────────────────────────────────────────────────────
    add_page_break(doc)
    doc.add_heading("Error handling", level=1)
    doc.add_paragraph(
        "Errors follow FastAPI's convention: a JSON object with a single 'detail' field (a string, or a "
        "structured list for 422 validation errors)."
    )
    add_code_block(doc, "{\n  \"detail\": \"Incorrect email or password\"\n}")
    doc.add_paragraph("Status codes used across the API:")
    make_table(doc, ["Status", "Meaning"], STATUS_CODES)

    # ── Endpoints ────────────────────────────────────────────────────────────
    for resource in RESOURCES:
        render_resource(doc, resource)

    # ── Payment integration notes ────────────────────────────────────────────
    add_page_break(doc)
    doc.add_heading("Payment integration notes", level=1)
    doc.add_paragraph(
        "Provider selection uses the Strategy pattern: a factory returns the Stripe or bKash strategy by "
        "PaymentProvider enum, and neither the routers nor the order logic branch on the provider — adding a "
        "new provider is a new strategy class only."
    )
    doc.add_heading("Idempotency", level=2)
    doc.add_paragraph(
        "payments.transaction_id is a UNIQUE column and the idempotency key for settlement. It is known at "
        "checkout time (the Stripe PaymentIntent id or the bKash paymentID). If a webhook/callback settles a "
        "payment that is already 'success', the operation is a no-op, and marking an order paid is likewise "
        "idempotent — so stock is never decremented twice. A duplicate checkout that would reuse a "
        "transaction_id is rejected with 409. Stock is reduced atomically with a WHERE stock >= quantity "
        "guard, which prevents overselling under concurrency."
    )
    doc.add_heading("Trust model", level=2)
    doc.add_paragraph(
        "Stripe webhooks are verified by signature (STRIPE_WEBHOOK_SECRET); only "
        "payment_intent.succeeded / payment_intent.payment_failed change state. bKash callbacks are trusted "
        "only after a server-side 'execute' call to bKash — the client-supplied 'status' is ignored. See "
        "docs/payment_flows.md for the full sequence diagrams."
    )

    # ── Data-model appendix ──────────────────────────────────────────────────
    add_page_break(doc)
    doc.add_heading("Appendix: Data models", level=1)
    doc.add_paragraph(
        "Field-level reference for every request and response schema. 'Required' applies to request bodies; "
        "response fields are always present (though some may be null)."
    )
    for name, fields in SCHEMA_APPENDIX:
        doc.add_heading(name, level=2)
        make_table(
            doc,
            ["Field", "Type", "Required", "Constraints / Default", "Description"],
            fields,
        )

    # ── Finalize ─────────────────────────────────────────────────────────────
    add_page_numbers(doc)
    enable_update_fields(doc)
    doc.save(str(OUTPUT))
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build()
